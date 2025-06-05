from docx import Document
import re
import os
from datetime import datetime, timedelta, date
from pathlib import Path

vandaag = datetime.today()
vorige_week = vandaag - timedelta(weeks=1)
weeknummer = vorige_week.isocalendar()[1]
weekmap = f"Debriefingsformulieren/Week{weeknummer}"

categorieen = [
    "OVERLAST PERSONEN",
    "JEUGDOVERLAST",
    "AFVALPROBLEMATIEK",
    "parkeeroverlast",
    "taken en opvallendheden"
]

resultaten = {cat: [] for cat in categorieen}


# Zoek de datum één keer, los van categorieën
for bestandsnaam in os.listdir(weekmap):
    if bestandsnaam.endswith(".docx"):
        pad = os.path.join(weekmap, bestandsnaam)
        doc = Document(pad)

        datum = None
        dienst = None

        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if "Datum dienst" in cell.text:
                        if i + 1 < len(row.cells):
                            datum = row.cells[i + 1].text.strip()
                        break
                    if "Soort dienst" in cell.text:
                        if i + 1 < len(row.cells):
                            dienst = row.cells[i + 1].text.strip()
                        break

        for table in doc.tables:
            rows = table.rows
            for i, row in enumerate(rows):
                rij_tekst = " ".join(cell.text.strip() for cell in row.cells)
                for cat in categorieen:
                    patroon = r"\b" + re.escape(cat) + r"\b"
                    if re.search(patroon, rij_tekst, re.IGNORECASE):
                        if i + 1 < len(rows):
                            tekst_volgende_rij = rows[i + 1].cells[0].text.strip()
                            if tekst_volgende_rij:  # alleen als er echt tekst is
                                resultaten[cat].append((datum, dienst, tekst_volgende_rij))

doc_out = Document()
doc_out.add_heading(f'Debriefingoverzicht Week {weeknummer}', 0)

def sorteersleutel(item):
    volgorde = {"ochtend": 0, "tussen": 1, "avond": 2}
    try:
        datum = datetime.strptime(item[0], "%d-%m-%Y")
    except Exception:
        datum = datetime.min

    dienst = item[1].lower()
    dienst_index = 99  # fallback als dienst onbekend is
    for dagdeel in volgorde:
        if dagdeel in dienst:
            dienst_index = volgorde[dagdeel]
            break

    return (datum, dienst_index)

for cat in resultaten:
    def parse_datum(item):
        try:
            return datetime.strptime(item[0], "%d-%m-%Y")
        except Exception:
            return datetime.min

    resultaten[cat].sort(key=sorteersleutel)


# for cat, items in resultaten.items():
#     print(f"\n🟦 {cat.upper()}")
#     if items:
#         # Filter items waarbij tekst niet leeg is
#         items_met_tekst = [item for item in items if item[2].strip() != ""]
#         if items_met_tekst:
#             for idx, (d, dienst, tekst) in enumerate(items_met_tekst, 1):
#                 print(f"\n🔹 {d} ({dienst}):\n{tekst}")


for cat, items in resultaten.items():
    if items:
        doc_out.add_heading(cat.upper(), level=1)
        for datum, dienst, tekst in items:
            doc_out.add_paragraph(f"{datum} ({dienst})", style='Heading 3')
            for regel in tekst.split('\n'):
                regel = regel.strip()
                if regel:
                    p = doc_out.add_paragraph(style='List Bullet')
                    p.add_run(regel)

output_pad = os.path.join(weekmap, f"Week {weeknummer} Debriefingsoverzicht.docx")
doc_out.save(output_pad)

print(f"✅ Document opgeslagen als: {output_pad}")
