import streamlit as st
from docx import Document
from datetime import datetime, timedelta
import re
import tempfile
import os
from io import BytesIO
import zipfile
from lxml import etree

# -------------------------------
# Inloggen met users in secrets
# -------------------------------
def login():
    st.title("🔐 Login")
    username = st.text_input("Gebruikersnaam")
    password = st.text_input("Wachtwoord", type="password")
    login_knop = st.button("Inloggen")

    if login_knop:
        if username in st.secrets["users"] and st.secrets["users"][username] == password:
            st.session_state["logged_in"] = True
            st.session_state["user"] = username
        else:
            st.error("Ongeldige gebruikersnaam of wachtwoord")

if "logged_in" not in st.session_state or not st.session_state["logged_in"]:
    login()
    st.stop()



def extract_date_picker_fields(docx_path):
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    with zipfile.ZipFile(docx_path) as docx_zip:
        with docx_zip.open('word/document.xml') as xml_file:
            tree = etree.parse(xml_file)
            root = tree.getroot()

            date_fields = []
            for sdt in root.findall('.//w:sdt', namespaces):
                # Kijk of dit een datumveld is
                if sdt.find('.//w:date', namespaces) is not None:
                    text_elements = sdt.findall('.//w:t', namespaces)
                    value = ''.join(t.text for t in text_elements if t.text)
                    date_fields.append(value)

            return date_fields


# -------------------------------
# Ingelogde content hieronder
# -------------------------------
st.title("📄 Debriefings Verwerker")

# Keuze voor onderdeel
onderdeel = st.radio("Kies onderdeel:", ["VOV", "Nieuw-West"])

# Categorieën per onderdeel
categorieen_NW = ["OVERLAST PERSONEN", "JEUGDOVERLAST", "AFVALPROBLEMATIEK", "parkeeroverlast", "taken en opvallendheden"]
categorieen_VOV = ["Jeugdoverlast", "Slapers/daklozen", "Geen/ongeldig vervoersbewijs", "Fietsen/steps/skaten/scooter", "Nooddeuren", "Roken", "Alcohol/drugs", "Diefstal", "Overig", "Werkopdracht 1", "Werkopdracht 2", "Werkopdracht 3", "Werkopdracht 4"]

categorieen = categorieen_VOV if onderdeel == "VOV" else categorieen_NW

uploaded_files = st.file_uploader(
    "Upload één of meerdere .docx-bestanden", 
    type=["docx"], 
    accept_multiple_files=True
)

if uploaded_files:
    resultaten = {cat: [] for cat in categorieen}

    for uploaded_file in uploaded_files:
        data = uploaded_file.read()
        file_stream = BytesIO(data)
        doc = Document(file_stream)

        datum = None
        dienst = None

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
            tmp_docx.write(data)
            tmp_docx_path = tmp_docx.name

        datumvelden = extract_date_picker_fields(tmp_docx_path)

        # Kies eventueel de eerste datum die je vindt
        if datumvelden:
            datum = datumvelden[0]  # Of loop door om specifieke te vinden


        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    label = row.cells[0].text.strip().lower().replace(":", "")
                    value = row.cells[1].text.strip()
                    if "datum dienst" in label and value:
                        datum = value
                    if "tijden + sector" in label and value:
                        dienst = value

        # Zoek categorieën en teksten
        for table in doc.tables:
            rows = table.rows
            for i, row in enumerate(rows):
                rij_tekst = " ".join(cell.text.strip() for cell in row.cells)
                for cat in categorieen:
                    patroon = r"\b" + re.escape(cat) + r"\b"
                    if re.search(patroon, rij_tekst, re.IGNORECASE):
                        if i + 1 < len(rows):
                            tekst_volgende_rij = rows[i + 1].cells[0].text.strip()
                            if tekst_volgende_rij:
                                resultaten[cat].append((datum, dienst, tekst_volgende_rij))

    def sorteerdagdelen(item):
        volgorde = {"ochtend": 0, "tussen": 1, "avond": 2}
        try:
            datum = datetime.strptime(item[0], "%d-%m-%Y")
        except Exception:
            datum = datetime.min

        dienst = item[1].lower() if item[1] else ""
        dienst_index = 99
        for dagdeel in volgorde:
            if dagdeel in dienst:
                dienst_index = volgorde[dagdeel]
                break
        return (datum, dienst_index)

    for cat in resultaten:
        resultaten[cat].sort(key=sorteerdagdelen)

    huidig_jaar = datetime.today().year
    huidige_week = datetime.today().isocalendar()[1]

    jaar_keuze = st.number_input("Selecteer jaar", min_value=2000, max_value=2100, value=huidig_jaar)
    week_keuze = st.number_input("Selecteer weeknummer", min_value=1, max_value=53, value=huidige_week - 1)

    weeknummer = int(week_keuze)
    jaar = int(jaar_keuze)

    doc_out = Document()
    doc_out.add_heading(f'Debriefingsoverzicht Week {weeknummer} - {jaar}', 0)

    for cat, items in resultaten.items():
        if items:
            doc_out.add_heading(cat.upper(), level=1)
            for datum, dienst, tekst in items:
                try:
                    datum_obj = datetime.strptime(datum, "%d-%m-%Y")
                    dag_van_week = datum_obj.strftime("%A")
                    dag_nl = {
                        "Monday": "Maandag",
                        "Tuesday": "Dinsdag",
                        "Wednesday": "Woensdag",
                        "Thursday": "Donderdag",
                        "Friday": "Vrijdag",
                        "Saturday": "Zaterdag",
                        "Sunday": "Zondag"
                    }.get(dag_van_week, dag_van_week)
                except Exception:
                    dag_nl = ""

                doc_out.add_paragraph(f"{dag_nl} {datum} ({dienst})", style='Heading 3')

                for regel in tekst.split('\n'):
                    regel = regel.strip()
                    if regel:
                        p = doc_out.add_paragraph(style='List Bullet')
                        p.add_run(regel)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_output:
        doc_out.save(tmp_output.name)
        tmp_output_path = tmp_output.name

    with open(tmp_output_path, "rb") as file:
        st.success("✅ Debriefing is gegenereerd!")
        st.download_button(
            label="📥 Download samenvatting",
            data=file,
            file_name=f"Week_{weeknummer}_Debriefingsoverzicht_{onderdeel}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
