import streamlit as st
from docx import Document
from datetime import datetime
import re
import tempfile
from io import BytesIO
from collections import defaultdict
from docx.shared import RGBColor

# -------------------------------
# Inloggen met users in secrets
# -------------------------------
def login():
    st.sidebar.title("🔐 Login")
    username = st.sidebar.text_input("Gebruikersnaam")
    password = st.sidebar.text_input("Wachtwoord", type="password")
    login_knop = st.sidebar.button("Inloggen")

    if login_knop:
        if username in st.secrets["users"] and st.secrets["users"][username] == password:
            st.session_state["logged_in"] = True
            st.session_state["user"] = username
        else:
            st.sidebar.error("Ongeldige gebruikersnaam of wachtwoord")

if "logged_in" not in st.session_state or not st.session_state["logged_in"]:
    login()
    st.stop()

# -------------------------------
# Ingelogde content hieronder
# -------------------------------
st.title("📄 Debriefings Verwerker")

categorieen = [
    "Vrijhouden van calamiteitenroutes",
    "In- en uitstroom van publiek",
    "Illegale evenementen in de openbare ruimte",
    "In hoeverre vielen andere vormen van overlast op",
    "Sfeerbeeld op straat",
    "Beschrijf hoe het publiek reageerde op de aanwezigheid van en contacten met THOR:",
    "Was er sprake van agressie en geweld (fysiek en//of verbaal) tegen collega's van THOR?",
    "Had je voldoende capaciteit om in te zetten?"
]

uploaded_files = st.file_uploader(
    "Upload één of meerdere .docx-bestanden", 
    type=["docx"], 
    accept_multiple_files=True
)

if uploaded_files:
    resultaten = {cat: [] for cat in categorieen}

    for uploaded_file in uploaded_files:
        data = uploaded_file.read()
        file_stream = BytesIO(data)
        doc = Document(file_stream)

        datum = None
        dienst = None
        inzetgebied = None

        # Zoek datum, dienst en inzetgebied
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if "Datum dienst" in cell.text and i + 1 < len(row.cells):
                        datum = row.cells[i + 1].text.strip()
                    if "Soort dienst" in cell.text and i + 1 < len(row.cells):
                        dienst = row.cells[i + 1].text.strip()
                    if "Inzetgebied" in cell.text and i + 1 < len(row.cells):
                        inzetgebied_raw = row.cells[i + 1].text.strip()
                        inzetgebied = "S105" if "S105" in inzetgebied_raw.upper() else inzetgebied_raw

        # Zoek categorieën en teksten
        for table in doc.tables:
            rows = table.rows
            for i, row in enumerate(rows):
                rij_tekst = " ".join(cell.text.strip() for cell in row.cells)
                for cat in categorieen:
                    patroon = r"\b" + re.escape(cat) + r"\b"
                    if re.search(patroon, rij_tekst, re.IGNORECASE):
                        if i + 1 < len(rows):
                            tekst_volgende_rij = rows[i + 1].cells[0].text.strip()
                            if tekst_volgende_rij:
                                resultaten[cat].append((datum, dienst, inzetgebied, tekst_volgende_rij))

    def sorteerdagdelen(item):
        volgorde = {"ochtend": 0, "tussen": 1, "avond": 2}
        try:
            datum_obj = datetime.strptime(item[0], "%d-%m-%Y")
        except Exception:
            datum_obj = datetime.min

        dienst = item[1].lower() if item[1] else ""
        dienst_index = 99
        for dagdeel in volgorde:
            if dagdeel in dienst:
                dienst_index = volgorde[dagdeel]
                break
        return (datum_obj, dienst_index)

    # Genereer document met structuur: Categorie -> Inzetgebied -> Dienst
    doc_out = Document()
    doc_out.add_heading(f'Debriefingsoverzicht Feest op de Ring', 0)

    for cat in categorieen:
        items = resultaten.get(cat, [])
        if not items:
            continue

        # Groepeer per inzetgebied binnen deze categorie
        inzetgebied_dict = defaultdict(list)
        for datum, dienst, inzetgebied, tekst in items:
            inzetgebied_dict[inzetgebied].append((datum, dienst, tekst))

        # Voeg rode kop toe voor de categorie
        p_cat = doc_out.add_paragraph()
        run = p_cat.add_run(cat.upper())
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Rood
        p_cat.style = 'Heading 1'

        for inzetgebied in sorted(inzetgebied_dict.keys()):
            doc_out.add_heading(f"{inzetgebied}", level=2)

            sorted_items = sorted(inzetgebied_dict[inzetgebied], key=lambda x: sorteerdagdelen((x[0], x[1])))

            for _, dienst, tekst in sorted_items:
                p_dienst = doc_out.add_paragraph(dienst)
                p_dienst.paragraph_format.space_after = 0
                run = p_dienst.runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                for regel in tekst.split('\n'):
                    regel = regel.strip()
                    if regel:
                        p = doc_out.add_paragraph(style='List Bullet')
                        p.add_run(regel)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_output:
        doc_out.save(tmp_output.name)
        tmp_output_path = tmp_output.name

    with open(tmp_output_path, "rb") as file:
        st.success("✅ Debriefing is gegenereerd!")
        st.download_button(
            label="📥 Download samenvatting",
            data=file,
            file_name=f"Debriefingsoverzicht.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

