import streamlit as st
from docx import Document
from datetime import datetime
import re
import tempfile
from io import BytesIO
from collections import defaultdict
from docx.shared import RGBColor

# -------------------------------
# Inloggen met users in secrets
# -------------------------------
# def login():
#     st.sidebar.title("🔐 Login")
#     username = st.sidebar.text_input("Gebruikersnaam")
#     password = st.sidebar.text_input("Wachtwoord", type="password")
#     login_knop = st.sidebar.button("Inloggen")

#     if login_knop:
#         if username in st.secrets["users"] and st.secrets["users"][username] == password:
#             st.session_state["logged_in"] = True
#             st.session_state["user"] = username
#         else:
#             st.sidebar.error("Ongeldige gebruikersnaam of wachtwoord")

# if "logged_in" not in st.session_state or not st.session_state["logged_in"]:
#     login()
#     st.stop()

# -------------------------------
# Ingelogde content hieronder
# -------------------------------
st.title("📄 Debriefings Verwerker")

categorieen = [
    "Vrijhouden van calamiteitenroutes en vaarroutes",
    "Toezien op in- en uitstroom van het evenement",
    "Illegale evenementen in de openbare ruimte",
    "In hoeverre vielen andere vormen van overlast op?",
    "Sfeerbeeld op straat",
    "Beschrijf hoe het publiek reageerde op de aanwezigheid van en contacten met THOR:",
    "Was er sprake van agressie en geweld (fysiek en/of verbaal) tegen collega's van THOR?"
]

uploaded_files = st.file_uploader(
    "Upload één of meerdere .docx-bestanden", 
    type=["docx"], 
    accept_multiple_files=True
)

def parse_nederlandse_datum(datum_str):
    maanden = {
        "januari": 1, "februari": 2, "maart": 3, "april": 4,
        "mei": 5, "juni": 6, "juli": 7, "augustus": 8,
        "september": 9, "oktober": 10, "november": 11, "december": 12
    }

    # Verwijder dagnaam (zoals "Vrijdag") als aanwezig
    delen = datum_str.strip().lower().split()
    if len(delen) >= 3:
        try:
            dag = int(delen[-3])
            maand = maanden.get(delen[-2])
            jaar = int(delen[-1])
            if maand:
                return datetime(jaar, maand, dag)
        except:
            pass
    return None

if uploaded_files:
    resultaten = {cat: [] for cat in categorieen}

    for uploaded_file in uploaded_files:
        data = uploaded_file.read()
        file_stream = BytesIO(data)
        doc = Document(file_stream)

        datum = None
        dienst = None
        inzetgebied = None

        # Zoek datum, dienst en inzetgebied
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if "Datum dienst" in cell.text and i + 1 < len(row.cells):
                        datum = row.cells[i + 1].text.strip()
                    if "Soort dienst" in cell.text and i + 1 < len(row.cells):
                        dienst = row.cells[i + 1].text.strip()
                    if "Inzetgebied" in cell.text and i + 1 < len(row.cells):
                        inzetgebied = row.cells[i + 1].text.strip()
                        

        # Zoek categorieën en teksten
        for table in doc.tables:
            rows = table.rows
            for i, row in enumerate(rows):
                rij_tekst = " ".join(cell.text.strip() for cell in row.cells)
                for cat in categorieen:
                    patroon = r"\b" + re.escape(cat) + r"\b"
                    if re.search(patroon, rij_tekst, re.IGNORECASE):
                        if i + 1 < len(rows):
                            tekst_volgende_rij = rows[i + 1].cells[0].text.strip()
                            if tekst_volgende_rij:
                                resultaten[cat].append((datum, dienst, inzetgebied, tekst_volgende_rij))

    def sorteerdagdelen(item):
        volgorde = {"ochtend": 0, "tussen": 1, "avond": 2}

        datum_obj = parse_nederlandse_datum(item[0]) or datetime.min

        dienst = item[1].lower() if item[1] else ""
        dienst_index = 99
        for dagdeel in volgorde:
            if dagdeel in dienst:
                dienst_index = volgorde[dagdeel]
                break

        return (datum_obj, dienst_index)


    # Genereer document met structuur: Datum -> Inzetgebied -> Categorie -> Dienst
    doc_out = Document()
    doc_out.add_heading(f'Debriefingsoverzicht Pride 2025', 0)

    # Verzamel en groepeer resultaten per datum > inzetgebied > categorie
    gestructureerd = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))

    for cat, items in resultaten.items():
        for datum, dienst, inzetgebied, tekst in items:
            gestructureerd[datum][inzetgebied][cat].append((dienst, tekst))

    # Sorteer op datum (chronologisch)
    gesorteerde_datums = sorted([d for d in gestructureerd.keys() if d], key=lambda d: parse_nederlandse_datum(d) or datetime.min)

    for datum in sorted(gestructureerd.keys(), key=lambda d: datetime.strptime(d, '%d-%m-%Y')):
        doc_out.add_heading(f"📅 {datum}", level=1)
        inzetgebieden = gestructureerd[datum]

        for inzetgebied in sorted([ig for ig in inzetgebieden.keys() if ig]):
            # Haal dienst op van de eerste categorie en eerste observatie
            categoriedata = inzetgebieden[inzetgebied]
            dienst_naam = None

            # Zoek eerste dienst van eerste categorie met observaties
            for cat in categorieen:
                if cat in categoriedata and categoriedata[cat]:
                    dienst_naam = categoriedata[cat][0][0]  # dienst uit (dienst, tekst)
                    break

            # Voeg inzetgebied + dienst samen in kop
            kop_tekst = f"📍 {inzetgebied}"
            if dienst_naam:
                kop_tekst += f" ({dienst_naam})"

            doc_out.add_heading(kop_tekst, level=2)

            for cat in categorieen:
                if cat not in categoriedata:
                    continue

                observaties = categoriedata[cat]
                if not observaties:
                    continue

                # Voeg rode categorie-kop toe
                p_cat = doc_out.add_paragraph()
                run = p_cat.add_run(cat.upper())
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                p_cat.style = 'Heading 3'

                # Sorteer observaties op dienst (tijd)
                sorted_obs = sorted(observaties, key=lambda x: sorteerdagdelen((datum, x[0])))

                for _, tekst in sorted_obs:
                    for regel in tekst.split('\n'):
                        regel = regel.strip()
                        if regel:
                            p = doc_out.add_paragraph(style='List Bullet')
                            p.add_run(regel)


    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_output:
        doc_out.save(tmp_output.name)
        tmp_output_path = tmp_output.name

    with open(tmp_output_path, "rb") as file:
        st.success("✅ Debriefing is gegenereerd!")
        st.download_button(
            label="📥 Download samenvatting",
            data=file,
            file_name=f"Debriefingsoverzicht.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # -----------------------------------------------
    # Tweede overzicht: Datum > Categorie > Inzetgebied
    # -----------------------------------------------
    doc_out2 = Document()
    doc_out2.add_heading(f'Debriefingsoverzicht Pride 2025 (per categorie)', 0)

    # Nieuwe structuur bouwen
    per_categorie = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    for cat, items in resultaten.items():
        for datum, dienst, inzetgebied, tekst in items:
            per_categorie[datum][cat][inzetgebied].append((dienst, tekst))

    # Chronologisch op datum
    for datum in sorted(per_categorie.keys(), key=lambda d: datetime.strptime(d, '%d-%m-%Y')):
        doc_out2.add_heading(f"📅 {datum}", level=1)
        categoriedata = per_categorie[datum]

        for cat in categorieen:
            if cat not in categoriedata:
                continue

            inzetgebieden = categoriedata[cat]
            if not inzetgebieden:
                continue

            # Voeg rode categorie-kop toe
            p_cat = doc_out2.add_paragraph()
            run = p_cat.add_run(cat.upper())
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            p_cat.style = 'Heading 2'

            for inzetgebied in sorted(inzetgebieden.keys()):
                observaties = inzetgebieden[inzetgebied]
                if not observaties:
                    continue

                # Zoek eerste dienst voor deze inzet
                dienst_naam = observaties[0][0] if observaties[0] else None
                inzet_kop = f"📍 {inzetgebied}"
                if dienst_naam:
                    inzet_kop += f" ({dienst_naam})"

                doc_out2.add_heading(inzet_kop, level=3)

                sorted_obs = sorted(observaties, key=lambda x: sorteerdagdelen((datum, x[0])))

                for _, tekst in sorted_obs:
                    for regel in tekst.split('\n'):
                        regel = regel.strip()
                        if regel:
                            p = doc_out2.add_paragraph(style='List Bullet')
                            p.add_run(regel)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_output2:
        doc_out2.save(tmp_output2.name)
        tmp_output_path2 = tmp_output2.name

    with open(tmp_output_path2, "rb") as file2:
        st.success("✅ Alternatief overzicht (per categorie) is gegenereerd!")
        st.download_button(
            label="📥 Download overzicht per categorie",
            data=file2,
            file_name=f"Debriefingsoverzicht_per_categorie.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


